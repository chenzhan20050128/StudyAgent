"""本项目内部使用的文档/网页解析工具。"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from typing import List

import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
from docx import Document


# ========= PDF / Word =========


@dataclass
class TextChunk:
    # 预留的数据结构：用于在需要“带来源的分块”时承载信息。
    # 当前主流程里实际分块由 LangChain splitter 完成，因此该结构更多是可扩展点。
    content: str
    source: str
    chunk_index: int
    char_count: int
    metadata: dict


class DocumentParser:
    """PDF / Word 文档解析为纯文本。

    这里只复用解析思路，不在这里做分块，分块交给上层的
    Splitter 保持全局一致策略。
    """

    def __init__(self, chunk_size: int = 800, overlap: int = 100) -> None:
        # 说明：该类当前只做“解析为纯文本”。
        # chunk_size/overlap 在现版本中并不参与实际切分（切分交给上层 TextSplitter），
        # 这里更像是历史遗留参数或未来扩展接口。
        self.chunk_size = chunk_size
        self.overlap = overlap

    def parse_pdf(self, path: str) -> str:
        """解析PDF文件为纯文本

        逻辑流程:
        1. 打开PDF文件并逐页遍历
        2. 从每页提取文本块(blocks)
        3. 过滤出纯文本块
        4. 为每页添加分页标记便于识别
        5. 将所有文本用换行符连接
        """
        # 业务目标：把 PDF 转成“尽可能干净的纯文本”，为后续分块/向量化做准备。
        # 技术策略：逐页提取 blocks，并只保留文本图层（过滤图片、矢量等非文本内容）。
        text_parts: List[str] = []
        # 使用上下文管理器打开PDF文件
        with fitz.open(path) as doc:
            # 遍历PDF的每一页，page_num是页码(0-indexed)
            for page_num, page in enumerate(doc):
                # 获取当前页的所有文本块，返回列表格式
                # blocks: (x0, y0, x1, y1, "text", block_no, block_type, ...)
                # 这里使用 block_type 过滤：只保留文本块
                blocks = page.get_text("blocks")
                # 从blocks中提取纯文本内容
                # b[4]是文本内容，b[6]==0表示是文本块(非图像等)
                page_text = [b[4] for b in blocks if b[6] == 0]
                # 添加页码分隔符，便于后续处理时识别页边界
                # 插入页分隔符：便于定位引用、也方便人类阅读与调试
                text_parts.append(f"--- 第{page_num + 1}页 ---")
                # 将当前页的所有文本块添加到总列表
                text_parts.extend(page_text)
        # 用换行符连接所有文本部分
        return "\n".join(text_parts)

    def parse_word(self, path: str) -> str:
        """解析Word文档为纯文本

        逻辑流程:
        1. 打开Word文件
        2. 提取所有段落文本(非空)
        3. 提取所有表格内容，按行组织
        4. 将所有内容用换行符连接
        """
        # 业务目标：尽可能完整保留 Word 的“正文段落 + 表格信息”。
        # 段落直接读取；表格按行拼接，保留单元格分隔符，避免结构信息完全丢失。
        # 使用 python-docx 库打开 Word 文档
        doc = Document(path)
        text_parts: List[str] = []

        # 提取文档中的所有段落
        for para in doc.paragraphs:
            # 只保留非空段落(去除纯空白行)
            if para.text.strip():
                text_parts.append(para.text)

        # 提取文档中的所有表格
        for table in doc.tables:
            # 添加表格标记
            text_parts.append("\n[表格]")
            # 逐行处理表格
            for row in table.rows:
                # 将同一行的单元格内容用" | "分隔符连接
                row_text = " | ".join(c.text.strip() for c in row.cells)
                text_parts.append(row_text)

        # 用换行符连接所有文本和表格内容
        return "\n".join(text_parts)

    def parse(self, path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            return self.parse_pdf(path)
        if ext in {".doc", ".docx"}:
            return self.parse_word(path)
        raise ValueError(f"不支持的格式: {ext}")


# ========= Web =========


class TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.text_parts: List[str] = []
        self.skip_content = False

    def handle_starttag(self, tag, attrs):  # type: ignore[override]
        # 过滤脚本/样式：网页正文抽取时，script/style 通常属于噪声
        if tag in ["script", "style"]:
            self.skip_content = True

    def handle_endtag(self, tag):  # type: ignore[override]
        # 对段落标签做换行控制：让输出更接近自然阅读的段落结构
        if tag in ["script", "style"]:
            self.skip_content = False
        elif tag in {"p", "br"}:
            self.text_parts.append("\n")

    def handle_data(self, data):  # type: ignore[override]
        if not self.skip_content:
            text = data.strip()
            if text:
                self.text_parts.append(text)

    def get_text(self) -> str:
        return "".join(self.text_parts)


class WebPageParser:
    """网页解析器，处理泛网页 + 微信 + 知乎。"""

    def __init__(self) -> None:
        # 统一请求头：降低被简单反爬策略拦截的概率
        self.headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 Chrome/91.0.4472.124 Safari/537.36"
            ),
            "Accept-Language": "zh-CN,zh;q=0.9",
        }

    def parse_weixin(self, url: str) -> str:
        # 微信文章：常见反爬更强，使用更像移动端的 UA + referer。
        # 抽取方式偏工程化：用正则提取标题与 <p> 段落，再去标签、做长度过滤。
        headers = self.headers.copy()
        headers.update(
            {
                "Referer": "https://mp.weixin.qq.com/",
                "User-Agent": (
                    "Mozilla/5.0 (Linux; Android 11; SM-G9810) "
                    "AppleWebKit/537.36 Chrome/88.0.4324.152 "
                    "Mobile Safari/537.36"
                ),
            }
        )
        resp = requests.get(url, headers=headers, timeout=30)
        resp.encoding = "utf-8"

        title_match = re.search(r'"msg_title":"([^"]*)"', resp.text)
        title = title_match.group(1) if title_match else "微信文章"

        paragraphs = re.findall(r"<p[^>]*>(.*?)</p>", resp.text, re.DOTALL)
        texts: List[str] = []
        for p in paragraphs:
            text = re.sub(r"<[^>]+>", "", p).strip()
            if text and len(text) > 10:
                texts.append(text)

        content = "\n".join(texts[:100])
        # 截断：避免正文过长导致后续 prompt 过大
        return f"【标题】{title}\n【URL】{url}\n\n{content}"

    def parse_zhihu(self, url: str) -> str:
        # 知乎：403 很常见，这里直接返回提示文本作为降级结果（不中断主流程）
        headers = self.headers.copy()
        headers.update(
            {
                "Referer": "https://www.zhihu.com/",
                "Accept": "text/html,application/xhtml+xml",
            }
        )
        resp = requests.get(url, headers=headers, timeout=30)
        if resp.status_code == 403:
            return (
                f"【标题】知乎文章\n【URL】{url}\n"
                "【提示】知乎防反爬虫较强，建议使用浏览器访问或利用官方 API"
            )

        resp.encoding = resp.apparent_encoding or "utf-8"
        soup = BeautifulSoup(resp.text, "html.parser")

        selectors = [
            ".Post-RichText",
            ".Post-content",
            "article",
            ".content",
        ]

        title = soup.find("h1")
        title_text = title.get_text(strip=True) if title else "知乎文章"

        content_text = ""
        for sel in selectors:
            elem = soup.select_one(sel)
            if elem is not None:
                content_text = elem.get_text(separator="\n", strip=True)
                break

        if not content_text:
            content_text = "（无法提取内容，知乎反爬虫限制）"

        return f"【标题】{title_text}\n【URL】{url}\n\n{content_text[:5000]}"

    def parse_generic(self, url: str) -> str:
        # 通用网页抽取：
        # 1) 去噪标签（script/style/nav/footer/header）
        # 2) 先尝试常见正文容器（article/main/...）
        # 3) 失败则退化为聚合较长的 <p> 段落
        resp = requests.get(url, headers=self.headers, timeout=30)
        if resp.encoding == "ISO-8859-1":
            resp.encoding = "utf-8"
        resp.encoding = resp.apparent_encoding or "utf-8"

        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        title = soup.find("h1") or soup.find("title")
        title_text = title.get_text(strip=True) if title else ""

        selectors = [
            "article",
            "main",
            ".content",
            ".post-content",
            ".entry-content",
            "#content",
            ".article-body",
        ]
        content_text = ""
        for sel in selectors:
            elem = soup.select_one(sel)
            if elem is not None:
                content_text = elem.get_text(separator="\n", strip=True)
                break

        if not content_text:
            paragraphs = soup.find_all("p")
            content_text = "\n\n".join(
                p.get_text(strip=True)
                for p in paragraphs
                if len(p.get_text().strip()) > 30
            )

        # 输出同样做截断，防止过长内容拖慢后续分块与 LLM 调用
        return f"【标题】{title_text}\n【URL】{url}\n\n{content_text[:5000]}"

    def parse(self, url: str) -> str:
        # 统一入口：按域名分发到具体策略；任何异常都降级为错误文本（避免阻断上传主流程）
        try:
            if "weixin.qq.com" in url or "mp.weixin.qq.com" in url:
                return self.parse_weixin(url)
            if "zhihu.com" in url:
                return self.parse_zhihu(url)
            return self.parse_generic(url)
        except Exception as exc:  # noqa: BLE001
            return f"【错误】{exc}"
